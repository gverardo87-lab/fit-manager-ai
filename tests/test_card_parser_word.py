"""
Test CardParser with a simulated Chiara-style Word document.

Creates a .docx file in memory with the same hybrid structure
as Chiara's real workout cards (mixed tables + free text + circuits).
"""
import io
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from docx import Document
from docx.shared import Pt
from core.card_parser import CardParser


def create_chiara_docx() -> bytes:
    """
    Create a .docx that mimics Chiara's real workout card format.

    Based on the actual document for "GIADA (MADAME)" pasted by the user.
    Structure:
    - Client name + measurement table
    - Weekly schedule table
    - "RISCALDAMENTO GENERALE" (free text bullets)
    - "RISCALDAMENTO SPECIFICO" (free text with circuit)
    - GIORNO 1: circuit with time-based work (table)
    - GIORNO 2: traditional exercises (table with SERIE-RIPETIZIONI)
    - GIORNO 3: PHA circuit (table)
    - Motivational text at the end
    """
    doc = Document()

    # Client name
    doc.add_paragraph("GIADA (MADAME)")

    # Measurement table (should be SKIPPED by parser)
    meas_table = doc.add_table(rows=2, cols=6)
    headers = ["PESO", "VITA", "FIANCHI", "COSCIA", "BRACCIO", "SENO"]
    for i, h in enumerate(headers):
        meas_table.rows[0].cells[i].text = h
    values = ["66", "74", "105", "60", "29", "102"]
    for i, v in enumerate(values):
        meas_table.rows[1].cells[i].text = v

    # Weekly schedule table (should be SKIPPED by parser)
    sched_table = doc.add_table(rows=2, cols=7)
    days = ["LUNEDI", "MARTEDI", "MERCOLEDI", "GIOVEDI", "VENERDI", "SABATO", "DOMENICA"]
    for i, d in enumerate(days):
        sched_table.rows[0].cells[i].text = d
    activities = ["ALLENAMENTO 1", "REST", "ALLENAMENTO 2", "REST", "ALLENAMENTO 3", "REST", "REST"]
    for i, a in enumerate(activities):
        sched_table.rows[1].cells[i].text = a

    # General warmup (free text - should mostly be skipped except exercises)
    doc.add_paragraph("RISCALDAMENTO GENERALE")
    doc.add_paragraph("Oltre all'allenamento 3 volte a settimana dovrai:")
    doc.add_paragraph("• 10 squat al mattino")
    doc.add_paragraph("• 5 burpee ogni sera")

    # Day 1 header
    doc.add_paragraph("GIORNO 1 – CIRCUITO")

    # Day 1 exercise table (circuit format with time-based work)
    t1 = doc.add_table(rows=6, cols=3)
    # Header
    t1.rows[0].cells[0].text = "ESERCIZIO"
    t1.rows[0].cells[1].text = "SERIE-RIPETIZIONI"
    t1.rows[0].cells[2].text = "RECUPERO"
    # Exercises
    t1.rows[1].cells[0].text = "THRUSTER"
    t1.rows[1].cells[1].text = '40"ON 20"OFF 3 GIRI'
    t1.rows[1].cells[2].text = "1'30\""

    t1.rows[2].cells[0].text = "AFFONDI POSTERIORI + ALZATE FRONTALI"
    t1.rows[2].cells[1].text = '40"ON 20"OFF 3 GIRI'
    t1.rows[2].cells[2].text = "1'30\""

    t1.rows[3].cells[0].text = "SUMO SQUAT + ROW"
    t1.rows[3].cells[1].text = '40"ON 20"OFF 3 GIRI'
    t1.rows[3].cells[2].text = "1'30\""

    t1.rows[4].cells[0].text = "CLEAN & JERK"
    t1.rows[4].cells[1].text = '40"ON 20"OFF 3 GIRI'
    t1.rows[4].cells[2].text = "1'30\""

    t1.rows[5].cells[0].text = "MOUNTAIN CLIMBER"
    t1.rows[5].cells[1].text = "20+20+20 4 GIRI"
    t1.rows[5].cells[2].text = "1'"

    # Day 2 header
    doc.add_paragraph("GIORNO 2 – FORZA")

    # Day 2 exercise table (traditional format)
    t2 = doc.add_table(rows=7, cols=4)
    t2.rows[0].cells[0].text = "ESERCIZIO"
    t2.rows[0].cells[1].text = "SERIE-RIPETIZIONI"
    t2.rows[0].cells[2].text = "SOVRACCARICHI"
    t2.rows[0].cells[3].text = "RECUPERO"

    t2.rows[1].cells[0].text = "HIP-THRUST"
    t2.rows[1].cells[1].text = "5x8/10"
    t2.rows[1].cells[2].text = "60kg"
    t2.rows[1].cells[3].text = "1'30\""

    t2.rows[2].cells[0].text = "STACCO DA TERRA"
    t2.rows[2].cells[1].text = "5x8/10"
    t2.rows[2].cells[2].text = "50kg"
    t2.rows[2].cells[3].text = "1'30\""

    t2.rows[3].cells[0].text = "SHOULDER PRESS"
    t2.rows[3].cells[1].text = "4x12"
    t2.rows[3].cells[2].text = "20kg"
    t2.rows[3].cells[3].text = "1'"

    t2.rows[4].cells[0].text = "PUSH DOWN"
    t2.rows[4].cells[1].text = "4x12"
    t2.rows[4].cells[2].text = "15kg"
    t2.rows[4].cells[3].text = "1'"

    t2.rows[5].cells[0].text = "LEG CURL"
    t2.rows[5].cells[1].text = "4x12 x gamba"
    t2.rows[5].cells[2].text = "20kg"
    t2.rows[5].cells[3].text = "1'"

    t2.rows[6].cells[0].text = "V-UP"
    t2.rows[6].cells[1].text = "20"
    t2.rows[6].cells[2].text = "-"
    t2.rows[6].cells[3].text = "-"

    # Day 3 header
    doc.add_paragraph("GIORNO 3 – PHA")

    # Day 3 table
    t3 = doc.add_table(rows=5, cols=3)
    t3.rows[0].cells[0].text = "ESERCIZIO"
    t3.rows[0].cells[1].text = "SERIE-RIPETIZIONI"
    t3.rows[0].cells[2].text = "RECUPERO"

    t3.rows[1].cells[0].text = "LAT MACHINE"
    t3.rows[1].cells[1].text = "4x12"
    t3.rows[1].cells[2].text = "1'"

    t3.rows[2].cells[0].text = "LEG PRESS"
    t3.rows[2].cells[1].text = "4x15"
    t3.rows[2].cells[2].text = "1'"

    t3.rows[3].cells[0].text = "FROG PUMP"
    t3.rows[3].cells[1].text = "20+20+20 3 GIRI"
    t3.rows[3].cells[2].text = "1'"

    t3.rows[4].cells[0].text = "HIP-THRUST MONOPODALICO"
    t3.rows[4].cells[1].text = "4x12 x gamba"
    t3.rows[4].cells[2].text = "1'30\""

    # Motivational text at the end (should be skipped)
    doc.add_paragraph("BUON LAVORO! Divertiti e non mollare mai! 💪")
    doc.add_paragraph("Ricordati che l'alimentazione è fondamentale per raggiungere i tuoi obiettivi.")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_chiara_word_parsing():
    """Test full parsing of Chiara-style .docx"""
    parser = CardParser()
    docx_bytes = create_chiara_docx()

    result = parser.parse_file(docx_bytes, "scheda_giada.docx")

    print("=" * 70)
    print("RISULTATO PARSING SCHEDA CHIARA")
    print("=" * 70)

    # Metadata
    print(f"\nMetadata:")
    print(f"  Goal: {result.metadata.detected_goal}")
    print(f"  Split: {result.metadata.detected_split}")
    print(f"  Sessions/week: {result.metadata.detected_sessions_per_week}")
    print(f"  Days found: {result.metadata.days_found}")
    print(f"  Client name: {result.metadata.client_name}")
    print(f"  Confidence: {result.parse_confidence}")

    # Exercises by day
    print(f"\nTotale esercizi trovati: {len(result.exercises)}")

    current_day = None
    for ex in result.exercises:
        if ex.day_section != current_day:
            current_day = ex.day_section
            print(f"\n--- {current_day or 'Warmup/No section'} ---")

        match_info = f"[{ex.canonical_id}:{ex.match_score}]" if ex.canonical_id else f"[NO MATCH:{ex.match_score}]"
        sets_info = f"Sets:{ex.sets}" if ex.sets else "Sets:-"
        reps_info = f"Reps:{ex.reps}" if ex.reps else "Reps:-"
        rest_info = f"Rest:{ex.rest_seconds}s" if ex.rest_seconds else ""
        load_info = f"Load:{ex.load_note}" if ex.load_note else ""

        print(f"  {ex.name:40s} {match_info:35s} {sets_info:10s} {reps_info:25s} {rest_info:10s} {load_info}")

    # Assertions
    print("\n" + "=" * 70)
    print("VERIFICHE")
    print("=" * 70)

    errors = []

    # Must find exercises
    if len(result.exercises) < 10:
        errors.append(f"FAIL: Solo {len(result.exercises)} esercizi (attesi >= 10)")
    else:
        print(f"OK: {len(result.exercises)} esercizi trovati (>= 10)")

    # Must detect 3 days
    if len(result.metadata.days_found) != 3:
        errors.append(f"FAIL: {len(result.metadata.days_found)} giorni trovati (attesi 3)")
    else:
        print(f"OK: 3 giorni trovati: {result.metadata.days_found}")

    # Must detect sessions
    if result.metadata.detected_sessions_per_week != 3:
        errors.append(f"FAIL: sessions={result.metadata.detected_sessions_per_week} (attesi 3)")
    else:
        print(f"OK: 3 sessioni/settimana rilevate")

    # Check key exercises are found
    exercise_names = [ex.name.lower() for ex in result.exercises]
    must_have = ["hip-thrust", "stacco da terra", "shoulder press", "lat machine", "leg press"]
    for name in must_have:
        if any(name in en for en in exercise_names):
            print(f"OK: '{name}' trovato")
        else:
            errors.append(f"FAIL: '{name}' NON trovato")

    # Check sets/reps parsing for key exercises
    for ex in result.exercises:
        if "hip-thrust" in ex.name.lower() and "monopodalico" not in ex.name.lower():
            if ex.sets == 5:
                print(f"OK: HIP-THRUST sets=5")
            else:
                errors.append(f"FAIL: HIP-THRUST sets={ex.sets} (attesi 5)")
            if ex.reps and "8" in ex.reps:
                print(f"OK: HIP-THRUST reps={ex.reps}")
            else:
                errors.append(f"FAIL: HIP-THRUST reps={ex.reps} (attesi 8/10)")
            if ex.load_note:
                print(f"OK: HIP-THRUST load={ex.load_note}")
            break

    # Check time-based circuit parsing (THRUSTER)
    for ex in result.exercises:
        if "thruster" in ex.name.lower():
            if ex.sets == 3:
                print(f"OK: THRUSTER sets=3 (giri)")
            else:
                errors.append(f"FAIL: THRUSTER sets={ex.sets} (attesi 3 giri)")
            if ex.reps and "on" in ex.reps.lower():
                print(f"OK: THRUSTER reps={ex.reps} (time-based)")
            else:
                errors.append(f"FAIL: THRUSTER reps={ex.reps} (attesi time-based)")
            if ex.rest_seconds == 90:
                print(f"OK: THRUSTER rest=90s")
            else:
                errors.append(f"FAIL: THRUSTER rest={ex.rest_seconds} (attesi 90)")
            break

    # Check measurement/schedule tables are NOT parsed as exercises
    meas_names = ["peso", "vita", "fianchi", "coscia", "braccio", "seno"]
    for name in meas_names:
        if name in exercise_names:
            errors.append(f"FAIL: '{name}' dal measurement table trovato come esercizio!")

    day_names = ["lunedi", "martedi", "mercoledi", "giovedi", "venerdi"]
    for name in day_names:
        if name in exercise_names:
            errors.append(f"FAIL: '{name}' dal schedule table trovato come esercizio!")

    if not any(name in exercise_names for name in meas_names + day_names):
        print(f"OK: Tabelle misure/schedule correttamente ignorate")

    # Confidence should be decent
    if result.parse_confidence >= 0.3:
        print(f"OK: Confidenza {result.parse_confidence} (>= 0.3)")
    else:
        errors.append(f"FAIL: Confidenza {result.parse_confidence} (attesa >= 0.3)")

    # Summary
    print("\n" + "=" * 70)
    if errors:
        print(f"RISULTATO: {len(errors)} ERRORI")
        for err in errors:
            print(f"  {err}")
    else:
        print("RISULTATO: TUTTI I TEST PASSATI!")
    print("=" * 70)

    return len(errors) == 0


if __name__ == "__main__":
    success = test_chiara_word_parsing()
    sys.exit(0 if success else 1)
